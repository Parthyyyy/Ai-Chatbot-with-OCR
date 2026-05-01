import os
import google.generativeai as ai
import cv2
import pytesseract
import pdfplumber
import docx
import pandas as pd
from PIL import Image
from PyPDF2 import PdfReader
from document_analyzer import DocumentAnalyzer

class EnhancedMultiFormatChatbot:
    def __init__(self, api_key):
        ai.configure(api_key=api_key)
        self.model = ai.GenerativeModel("gemini-2.5-flash")
        self.content = ""
        self.chat = None
        self.current_file = ""
        self.extracted_tables = []
        self.document_analyzer = DocumentAnalyzer()

    def set_context(self, text):
        self.context = text

    def load_combined_text(self, combined_text):
        """Load multiple PDF contents together into a single chat context"""
        self.text_context = combined_text

        context_prompt = f"""
        You are an AI assistant that answers questions based on document content, including structured data like tables and forms.

    IMPORTANT INSTRUCTIONS:
    - Give direct, accurate answers based on the extracted content
    - Pay special attention to tabular data and structured information
    - When asked about educational qualifications, look for institution names, years, and percentages
    - For B.Tech questions, specifically look for university/college names in the education section
    - Be specific with names, dates, and numbers when available
    - If information is in a table format, extract the exact values

    DOCUMENT CONTENT:
        {self.text_context}

    The document has been processed with enhanced table extraction. Answer questions accurately based on this content.
    """

        try:
            self.chat = self.model.start_chat()
            response = self.chat.send_message(context_prompt)
            print("✅ Combined documents loaded successfully!")
            return True
        except Exception as e:
            print(f"[ERROR] Failed to load combined content into chat: {e}")
            self.chat = None
        return False

    def preprocess_image_for_ocr(self, image_path):
        """Preprocess image to improve OCR accuracy"""
        try:
            # Read image using OpenCV
            img = cv2.imread(image_path)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply noise reduction
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Apply adaptive thresholding to get better contrast
            thresh = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Morphological operations to clean up the image
            kernel = np.ones((2, 2), np.uint8)
            processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            return processed
            
        except Exception as e:
            print(f"Error in image preprocessing: {e}")
            return cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    
    def detect_and_extract_tables(self, image_path):
        """Enhanced table detection and extraction"""
        try:
            img = cv2.imread(image_path)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply threshold to get binary image
            thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            
            # Detect horizontal lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
            horizontal_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
            
            # Detect vertical lines
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
            vertical_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
            
            # Combine horizontal and vertical lines
            table_mask = cv2.addWeighted(horizontal_lines, 0.5, vertical_lines, 0.5, 0.0)
            
            # Find contours to detect table cells
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Sort contours by area and keep only significant ones
            contours = sorted(contours, key=cv2.contourArea, reverse=True)
            significant_contours = [c for c in contours if cv2.contourArea(c) > 100]
            
            return len(significant_contours) > 3, significant_contours
            
        except Exception as e:
            print(f"Error in table detection: {e}")
            return False, []
    
    def extract_table_with_coordinates(self, image_path):
        """Extract table data using coordinate-based approach"""
        try:
            image = Image.open(image_path)
            
            # Get OCR data with coordinates
            ocr_data = pytesseract.image_to_data(
                image, 
                config='--psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,()-%/ ',
                output_type=pytesseract.Output.DICT
            )
            
            # Filter confident detections
            confident_data = []
            for i in range(len(ocr_data['text'])):
                if int(ocr_data['conf'][i]) > 40:  # Lower threshold for better detection
                    text = ocr_data['text'][i].strip()
                    if text and len(text) > 1:  # Ignore single characters
                        confident_data.append({
                            'text': text,
                            'left': ocr_data['left'][i],
                            'top': ocr_data['top'][i],
                            'width': ocr_data['width'][i],
                            'height': ocr_data['height'][i]
                        })
            
            if not confident_data:
                return None
            
            # Group data into rows based on Y coordinates
            rows = {}
            y_tolerance = 15  # Pixels tolerance for same row
            
            for item in confident_data:
                y = item['top']
                
                # Find existing row or create new one
                row_key = None
                for existing_y in rows.keys():
                    if abs(y - existing_y) <= y_tolerance:
                        row_key = existing_y
                        break
                
                if row_key is None:
                    row_key = y
                    rows[row_key] = []
                
                rows[row_key].append(item)
            
            # Sort rows by Y coordinate and create table
            sorted_rows = sorted(rows.keys())
            table_data = []
            
            for row_y in sorted_rows:
                # Sort items in row by X coordinate
                row_items = sorted(rows[row_y], key=lambda x: x['left'])
                row_text = [item['text'] for item in row_items]
                
                if len(row_text) > 1:  # Only include rows with multiple columns
                    table_data.append(row_text)
            
            return table_data
            
        except Exception as e:
            print(f"Error in coordinate-based extraction: {e}")
            return None
    
    def format_extracted_table(self, table_data):
        """Format extracted table data into readable text"""
        if not table_data:
            return ""
        
        try:
            # Create DataFrame
            if len(table_data) > 1:
                # First row as headers if it looks like headers
                headers = table_data[0]
                data_rows = table_data[1:]
                
                # Pad rows to match header length
                max_cols = len(headers)
                for i, row in enumerate(data_rows):
                    while len(row) < max_cols:
                        row.append("")
                    data_rows[i] = row[:max_cols]  # Trim if too long
                
                df = pd.DataFrame(data_rows, columns=headers)
                formatted_table = df.to_string(index=False)
                
                # Also create a structured text representation
                structured_text = f"\n=== TABLE STRUCTURE ===\n"
                for _, row in df.iterrows():
                    for col, val in row.items():
                        if val.strip():
                            structured_text += f"{col}: {val}\n"
                    structured_text += "---\n"
                
                return f"{formatted_table}\n{structured_text}"
            else:
                return str(table_data)
                
        except Exception as e:
            print(f"Error formatting table: {e}")
            return str(table_data)
    
    def extract_text_with_enhanced_table_support(self, image_path):
        """Enhanced text extraction with better table and section support"""
        try:
            print("Starting enhanced image extraction with section analysis...")
            
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
                
            # Get detailed OCR data with bounding boxes
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config='--psm 6')
            
            # Process text elements
            text_elements = []
            for i in range(len(ocr_data['text'])):
                text = ocr_data['text'][i].strip()
                conf = int(ocr_data['conf'][i])
                if text and conf > 50:
                     text_elements.append({
                         'text': text,
                         'left': ocr_data['left'][i],
                         'top': ocr_data['top'][i],
                         'width': ocr_data['width'][i],
                         'height': ocr_data['height'][i],
                         'line_num': ocr_data['line_num'][i],
                         'block_num': ocr_data['block_num'][i]
                     })
            
            # Sort elements by position
            text_elements.sort(key=lambda x: (x['top'], x['left']))
            
            # Reconstruct lines
            lines = []
            current_line = []
            y_tolerance = 10
            
            for element in text_elements:
                if not current_line:
                    current_line.append(element)
                else:
                    avg_line_height = sum(el['height'] for el in current_line) / len(current_line)
                    if abs(element['top'] - current_line[-1]['top']) < y_tolerance or \
                       (element['top'] + element['height'] > current_line[-1]['top'] and element['top'] < current_line[-1]['top'] + current_line[-1]['height']):
                         current_line.append(element)
                    else:
                        lines.append(current_line)
                        current_line = [element]
            
            if current_line:
                lines.append(current_line)
                
            # Process lines
            processed_lines = []
            for line_elements in lines:
                 line_elements.sort(key=lambda x: x['left'])
                 processed_lines.append(" ".join([el['text'] for el in line_elements]))

            # Analyze document content
            text_content = "\n".join(processed_lines)
            analysis = self.document_analyzer.analyze_document(text_content, image_path)
            
            # Format the analysis results
            structured_text = self.document_analyzer.format_analysis(analysis)
            
            # Add raw OCR outputs for redundancy
            all_ocr_outputs = ""
            all_ocr_outputs += "\n=== RAW OCR (PSM 6) ===\n" + pytesseract.image_to_string(image, config='--psm 6') + "\n\n"
            all_ocr_outputs += "=== RAW OCR (PSM 4) ===\n" + pytesseract.image_to_string(image, config='--psm 4') + "\n\n"
            all_ocr_outputs += "=== RAW OCR (PSM 11) ===\n" + pytesseract.image_to_string(image, config='--psm 11') + "\n\n"

            structured_text += all_ocr_outputs
            
            return structured_text.strip()
            
        except Exception as e:
            print(f"Error in enhanced image extraction: {e}")
            return self.extract_text_from_image_standard(image_path)
    
    def extract_email_addresses(self, text):
        """Extract email addresses from text using improved extraction"""
        return extract_email_addresses_improved(text, self.current_image_path if hasattr(self, 'current_image_path') else None)
    
    def extract_education_patterns(self, text):
        """Extract education-related patterns from text, including percentages"""
        try:
            education_info = []
            
            # Look for B.Tech patterns
            btech_patterns = [
                r'B\.?Tech\.?\s*(?:\([^)]*\))?\s*[:\-]?\s*([A-Za-z\s&]+(?:College|University|Institute))',
                r'B\.?Tech\.?\s*(?:\([^)]*\))?\s*([A-Za-z\s&]+(?:College|University|Institute))',
                r'(?:B\.?Tech\.?.*?)([A-Za-z\s&]+(?:Engineering|College|University|Institute))',
            ]
            
            for pattern in btech_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    institution = match.group(1).strip()
                    if len(institution) > 3:  # Filter out too short matches
                        education_info.append(f"B.Tech Institution: {institution}")
            
            # Look for year patterns near B.Tech
            year_patterns = r'B\.?Tech\.?.*?(?:20\d{2}|19\d{2})'
            year_matches = re.finditer(year_patterns, text, re.IGNORECASE)
            for match in year_matches:
                education_info.append(f"B.Tech Year Info: {match.group()}")
            
            # Look for percentage patterns specifically near education keywords (refined)
            percentage_patterns_education = r'(?:B\.?Tech\.?|Diploma|H\.?S\.?C\.?E\.?|S\.?S\.?C\.?E\.?).*?(\d{1,2}(?:.\d{1,2})?)\s*%?' # Look for 1-2 digits, optional decimal, optional %
            percentage_matches_education = re.finditer(percentage_patterns_education, text, re.IGNORECASE)
            for match in percentage_matches_education:
                percentage = match.group(1).strip()
                if percentage:
                    # Basic validation: ensure percentage is reasonable (e.g., <= 100)
                    try:
                        percentage_value = float(percentage)
                        if 0 <= percentage_value <= 100:
                            education_info.append(f"Education Percentage Found (Refined): {percentage}%")
                        else:
                            print(f"Warning: Extracted percentage {percentage}% seems unreasonable.")
                    except ValueError:
                        print(f"Warning: Could not convert extracted percentage {percentage} to float.")
            
            # Broader percentage pattern as a fallback (also refined)
            percentage_patterns_general = r'\b(\d{1,2}(?:.\d{1,2})?)\s*%?' # Look for 1-2 digits, optional decimal, optional %
            percentage_matches_general = re.finditer(percentage_patterns_general, text)
            for match in percentage_matches_general:
                percentage = match.group(1).strip()
                if percentage:
                    # Basic validation: ensure percentage is reasonable (e.g., <= 100)
                    try:
                        percentage_value = float(percentage)
                        if 0 <= percentage_value <= 100 and f"Education Percentage Found (Refined): {percentage}%" not in education_info:
                            education_info.append(f"Percentage Found (General, Refined): {percentage}%")
                    except ValueError:
                        print(f"Warning: Could not convert extracted percentage {percentage} to float (General).")
            
            return '\n'.join(education_info) if education_info else ""
            
        except Exception as e:
            print(f"Error in pattern extraction: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text using pdfplumber for PDF files with table support"""
        try:
            text = ""
            tables_text = ""
            
            with pdfplumber.open(pdf_path) as pdf:
                print(f"Processing {len(pdf.pages)} pages...")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract regular text
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        tables_text += f"\n--- Tables from Page {page_num} ---\n"
                        for table_num, table in enumerate(tables, 1):
                            tables_text += f"\nTable {table_num}:\n"
                            # Convert table to readable format
                            if table:
                                df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                                tables_text += df.to_string(index=False) + "\n"
                    
                    print(f"Processed page {page_num}")
            
            # Combine regular text and tables
            combined_text = text
            if tables_text:
                combined_text += "\n\n=== EXTRACTED TABLES ===\n" + tables_text
                
            return combined_text.strip()
        
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None
    
    def extract_text_from_image_standard(self, image_path):
        """Standard image text extraction"""
        try:
            print("Using standard OCR extraction...")
            
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            extracted_text = pytesseract.image_to_string(image, config='--psm 6')
            
            if extracted_text.strip():
                return extracted_text.strip()
            else:
                return "No readable text found in the image."
                
        except Exception as e:
            print(f"Error reading image: {e}")
            return None
    
    def extract_text_from_image(self, image_path):
        """Main image text extraction method"""
        return self.extract_text_with_enhanced_table_support(image_path)
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from DOC/DOCX files"""
        try:
            text = ""
            doc = docx.Document(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables from docx
            tables_text = ""
            for table_num, table in enumerate(doc.tables, 1):
                tables_text += f"\nTable {table_num}:\n"
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data[0] else None)
                    tables_text += df.to_string(index=False) + "\n"
            
            combined_text = text
            if tables_text:
                combined_text += "\n\n=== EXTRACTED TABLES ===\n" + tables_text
                
            return combined_text.strip()
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None
    
    def load_file(self, file_path):
        """Load PDF, DOC/DOCX or image file and prepare for Q&A"""
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found!")
            return False
        
        # Determine file type
       
        file_extension = os.path.splitext(file_path.lower())[1]
        
        print(f"Loading file: {os.path.basename(file_path)}")
        
        if file_extension == '.pdf':
            self.content = self.extract_text_from_pdf(file_path)
            file_type = "PDF"
        elif file_extension in ['.doc', '.docx']:
            self.content = self.extract_text_from_docx(file_path)
            file_type = "DOC/DOCX"
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
            self.content = self.extract_text_from_image(file_path)
            file_type = "Image"
        else:
            print(f"Unsupported file format: {file_extension}")
            print("Supported formats: PDF, DOC, DOCX, JPG, JPEG, PNG, BMP, TIFF, WEBP")
            return False
        
        if not self.content:
            print(f"Failed to extract text from {file_type}!")
            return False
        
        # Show preview of extracted content for debugging
        print("\n=== EXTRACTED CONTENT PREVIEW ===")
        print(self.content[:1000] + "..." if len(self.content) > 1000 else self.content)
        print("=== END PREVIEW ===\n")
        
        # Truncate content if too long
        if len(self.content) > 50000:
            self.content = self.content[:50000] + "\n[Content truncated due to length...]"
            print("Note: Content was truncated due to length limits.")
        
        self.current_file = os.path.basename(file_path)
        
        # Enhanced context prompt for better table understanding
        context_prompt = f"""
        You are an AI assistant that answers questions based on document content, including structured data like tables and forms.
        
        IMPORTANT INSTRUCTIONS:
        - Give direct, accurate answers based on the extracted content
        - Pay special attention to tabular data and structured information
        - When asked about educational qualifications, look for institution names, years, and percentages
        - For B.Tech questions, specifically look for university/college names in the education section
        - Be specific with names, dates, and numbers when available
        - If information is in a table format, extract the exact values
        
        DOCUMENT CONTENT:
        {self.content}
        
        The document has been processed with enhanced table extraction. Answer questions accurately based on this content.
        """
        
        try:
            self.chat = self.model.start_chat()
            response = self.chat.send_message(context_prompt)
            print(f"✅ {file_type} loaded successfully! Enhanced table extraction completed.")
            return True
        except Exception as e:
            print(f"Error initializing chat: {e}")
            return False
    
    def ask_question(self, question):
        """Ask a question with enhanced context understanding"""
        if not self.chat:
            return "Please load a file first using 'load <file_path>' command."
        
        try:
            # Enhanced question processing
            base_question_prompt = f"""
Question: {question}

Please provide a direct, specific answer based on the document content. The document has been analyzed for:
- Education details (degrees, institutions, years, scores)
- Contact information (email, phone, address)
- Skills and qualifications
- Professional experience
- Projects and achievements
- Other relevant sections

Please focus on the specific information requested in the question.
"""

            response = self.chat.send_message(base_question_prompt)
            return response.text

        except Exception as e:
            return f"Error: {e}"
    
    def show_content_preview(self):
        """Show a preview of the extracted content"""
        if not self.content:
            print("No content loaded.")
            return
        
        preview = self.content[:1000] + "..." if len(self.content) > 1000 else self.content
        print(f"\n--- Content Preview from {self.current_file} ---")
        print(preview)
        print("--- End Preview ---\n")
