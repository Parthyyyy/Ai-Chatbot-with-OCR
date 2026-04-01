import google.generativeai as ai
import pdfplumber
import os
import ollama
from PIL import Image
import pytesseract
from flask import Flask, render_template, request, jsonify, send_from_directory, session, request, redirect
import uuid
from dotenv import load_dotenv
import docx
import boto3
import google.generativeai as genai
import cv2
import numpy as np
import pandas as pd
from tabulate import tabulate
import re
from datetime import datetime
from email_validator import clean_email, is_valid_email_format, extract_email_addresses_improved
from document_analyzer import DocumentAnalyzer
import threading
from PyPDF2 import PdfReader

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.secret_key = os.getenv('SECRET_KEY', 'add8e6b193e17b39dd36ab85216f2c4c')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOG_DIR = os.path.join(BASE_DIR, "logs")
CHAT_LOG_PATH = os.path.join(LOG_DIR, "chat_log.txt")

def log_conversation(user_input, bot_response, override=False):
    try:
        log_path = session.get('chat_log_path')
        if not log_path:
            print("⚠️ No log path set for this session.")
            return
        full_path = os.path.join(os.getcwd(), log_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        mode = 'w' if override else 'a'
        with open(full_path, mode, encoding='utf-8') as f:
            f.write(f"User: {user_input}\n")
            f.write(f"Bot: {bot_response}\n\n")
        print("📝 Logged to:", full_path)
    except Exception as e:
        print("❌ Logging error:",e)


        os.makedirs(LOG_DIR, exist_ok=True)
        mode = 'w' if override else 'a'
        with open(CHAT_LOG_PATH, mode, encoding='utf-8') as f:
            f.write(f"User: {user_input}\n")
            f.write(f"Bot: {bot_response}\n\n")
        print("📝 Chat logged successfully at ", CHAT_LOG_PATH)
    except Exception as e:
        print("❌ Error logging chat:", e)

ALLOWED_EXTENSIONS = {'pdf', 'txt', 'doc', 'docx', 'rtf'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_polly_audio(text, filename="response.mp3", voice="Matthew"):

    static_path=os.path.join('Ai_chatbot', 'static')
    os.makedirs(static_path, exist_ok=True)

    output_path=os.path.join(static_path,filename)
    try :
        polly=boto3.client('polly', region_name='us-east-1')

        response=polly.synthesize_speech(
            Text=text,
            OutputFormat='mp3',
            VoiceId=voice
        )

        print("[Polly RAW RESPONSE]", response)
        if 'AudioStream' not in response:
            print("[POLLY ERROR] AudioStream not in response")
            if 'error' in response:
                print("[Polly ERROR DETAILS]:", response['Error'])
        
        with open(output_path,'wb') as f:
            f.write(response['AudioStream'].read())
        return f"/static/{filename}"
    
    except Exception as e:
        print("[Polly ERROR]", str(e))
        return None


def extract_text_from_pdf(path):
        text = ""
        with open(path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ''
        return text

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

class EnhancedMultiFormatChatbot:
    def __init__(self, api_key):
        ai.configure(api_key=api_key)
        self.model = ai.GenerativeModel("gemini-2.5-flash")
        self.content = ""
        self.chat = None
        self.current_file = ""
        self.extracted_tables = []
        self.document_analyzer = DocumentAnalyzer()

    def set_context(self, text):
        self.context = text

    def load_combined_text(self, combined_text):
        """Load multiple PDF contents together into a single chat context"""
        self.text_context = combined_text

        context_prompt = f"""
        You are an AI assistant that answers questions based on document content, including structured data like tables and forms.

    IMPORTANT INSTRUCTIONS:
    - Give direct, accurate answers based on the extracted content
    - Pay special attention to tabular data and structured information
    - When asked about educational qualifications, look for institution names, years, and percentages
    - For B.Tech questions, specifically look for university/college names in the education section
    - Be specific with names, dates, and numbers when available
    - If information is in a table format, extract the exact values

    DOCUMENT CONTENT:
        {self.text_context}

    The document has been processed with enhanced table extraction. Answer questions accurately based on this content.
    """

        try:
            self.chat = self.model.start_chat()
            response = self.chat.send_message(context_prompt)
            print("✅ Combined documents loaded successfully!")
            return True
        except Exception as e:
            print(f"[ERROR] Failed to load combined content into chat: {e}")
            self.chat = None
        return False

    def preprocess_image_for_ocr(self, image_path):
        """Preprocess image to improve OCR accuracy"""
        try:
            # Read image using OpenCV
            img = cv2.imread(image_path)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply noise reduction
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Apply adaptive thresholding to get better contrast
            thresh = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Morphological operations to clean up the image
            kernel = np.ones((2, 2), np.uint8)
            processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            return processed
            
        except Exception as e:
            print(f"Error in image preprocessing: {e}")
            return cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    
    def detect_and_extract_tables(self, image_path):
        """Enhanced table detection and extraction"""
        try:
            img = cv2.imread(image_path)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply threshold to get binary image
            thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            
            # Detect horizontal lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
            horizontal_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
            
            # Detect vertical lines
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
            vertical_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
            
            # Combine horizontal and vertical lines
            table_mask = cv2.addWeighted(horizontal_lines, 0.5, vertical_lines, 0.5, 0.0)
            
            # Find contours to detect table cells
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Sort contours by area and keep only significant ones
            contours = sorted(contours, key=cv2.contourArea, reverse=True)
            significant_contours = [c for c in contours if cv2.contourArea(c) > 100]
            
            return len(significant_contours) > 3, significant_contours
            
        except Exception as e:
            print(f"Error in table detection: {e}")
            return False, []
    
    def extract_table_with_coordinates(self, image_path):
        """Extract table data using coordinate-based approach"""
        try:
            image = Image.open(image_path)
            
            # Get OCR data with coordinates
            ocr_data = pytesseract.image_to_data(
                image, 
                config='--psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,()-%/ ',
                output_type=pytesseract.Output.DICT
            )
            
            # Filter confident detections
            confident_data = []
            for i in range(len(ocr_data['text'])):
                if int(ocr_data['conf'][i]) > 40:  # Lower threshold for better detection
                    text = ocr_data['text'][i].strip()
                    if text and len(text) > 1:  # Ignore single characters
                        confident_data.append({
                            'text': text,
                            'left': ocr_data['left'][i],
                            'top': ocr_data['top'][i],
                            'width': ocr_data['width'][i],
                            'height': ocr_data['height'][i]
                        })
            
            if not confident_data:
                return None
            
            # Group data into rows based on Y coordinates
            rows = {}
            y_tolerance = 15  # Pixels tolerance for same row
            
            for item in confident_data:
                y = item['top']
                
                # Find existing row or create new one
                row_key = None
                for existing_y in rows.keys():
                    if abs(y - existing_y) <= y_tolerance:
                        row_key = existing_y
                        break
                
                if row_key is None:
                    row_key = y
                    rows[row_key] = []
                
                rows[row_key].append(item)
            
            # Sort rows by Y coordinate and create table
            sorted_rows = sorted(rows.keys())
            table_data = []
            
            for row_y in sorted_rows:
                # Sort items in row by X coordinate
                row_items = sorted(rows[row_y], key=lambda x: x['left'])
                row_text = [item['text'] for item in row_items]
                
                if len(row_text) > 1:  # Only include rows with multiple columns
                    table_data.append(row_text)
            
            return table_data
            
        except Exception as e:
            print(f"Error in coordinate-based extraction: {e}")
            return None
    
    def format_extracted_table(self, table_data):
        """Format extracted table data into readable text"""
        if not table_data:
            return ""
        
        try:
            # Create DataFrame
            if len(table_data) > 1:
                # First row as headers if it looks like headers
                headers = table_data[0]
                data_rows = table_data[1:]
                
                # Pad rows to match header length
                max_cols = len(headers)
                for i, row in enumerate(data_rows):
                    while len(row) < max_cols:
                        row.append("")
                    data_rows[i] = row[:max_cols]  # Trim if too long
                
                df = pd.DataFrame(data_rows, columns=headers)
                formatted_table = df.to_string(index=False)
                
                # Also create a structured text representation
                structured_text = f"\n=== TABLE STRUCTURE ===\n"
                for _, row in df.iterrows():
                    for col, val in row.items():
                        if val.strip():
                            structured_text += f"{col}: {val}\n"
                    structured_text += "---\n"
                
                return f"{formatted_table}\n{structured_text}"
            else:
                return str(table_data)
                
        except Exception as e:
            print(f"Error formatting table: {e}")
            return str(table_data)
    
    def extract_text_with_enhanced_table_support(self, image_path):
        """Enhanced text extraction with better table and section support"""
        try:
            print("Starting enhanced image extraction with section analysis...")
            
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
                
            # Get detailed OCR data with bounding boxes
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config='--psm 6')
            
            # Process text elements
            text_elements = []
            for i in range(len(ocr_data['text'])):
                text = ocr_data['text'][i].strip()
                conf = int(ocr_data['conf'][i])
                if text and conf > 50:
                     text_elements.append({
                         'text': text,
                         'left': ocr_data['left'][i],
                         'top': ocr_data['top'][i],
                         'width': ocr_data['width'][i],
                         'height': ocr_data['height'][i],
                         'line_num': ocr_data['line_num'][i],
                         'block_num': ocr_data['block_num'][i]
                     })
            
            # Sort elements by position
            text_elements.sort(key=lambda x: (x['top'], x['left']))
            
            # Reconstruct lines
            lines = []
            current_line = []
            y_tolerance = 10
            
            for element in text_elements:
                if not current_line:
                    current_line.append(element)
                else:
                    avg_line_height = sum(el['height'] for el in current_line) / len(current_line)
                    if abs(element['top'] - current_line[-1]['top']) < y_tolerance or \
                       (element['top'] + element['height'] > current_line[-1]['top'] and element['top'] < current_line[-1]['top'] + current_line[-1]['height']):
                         current_line.append(element)
                    else:
                        lines.append(current_line)
                        current_line = [element]
            
            if current_line:
                lines.append(current_line)
                
            # Process lines
            processed_lines = []
            for line_elements in lines:
                 line_elements.sort(key=lambda x: x['left'])
                 processed_lines.append(" ".join([el['text'] for el in line_elements]))

            # Analyze document content
            text_content = "\n".join(processed_lines)
            analysis = self.document_analyzer.analyze_document(text_content, image_path)
            
            # Format the analysis results
            structured_text = self.document_analyzer.format_analysis(analysis)
            
            # Add raw OCR outputs for redundancy
            all_ocr_outputs = ""
            all_ocr_outputs += "\n=== RAW OCR (PSM 6) ===\n" + pytesseract.image_to_string(image, config='--psm 6') + "\n\n"
            all_ocr_outputs += "=== RAW OCR (PSM 4) ===\n" + pytesseract.image_to_string(image, config='--psm 4') + "\n\n"
            all_ocr_outputs += "=== RAW OCR (PSM 11) ===\n" + pytesseract.image_to_string(image, config='--psm 11') + "\n\n"

            structured_text += all_ocr_outputs
            
            return structured_text.strip()
            
        except Exception as e:
            print(f"Error in enhanced image extraction: {e}")
            return self.extract_text_from_image_standard(image_path)
    
    def extract_email_addresses(self, text):
        """Extract email addresses from text using improved extraction"""
        return extract_email_addresses_improved(text, self.current_image_path if hasattr(self, 'current_image_path') else None)
    
    def extract_education_patterns(self, text):
        """Extract education-related patterns from text, including percentages"""
        try:
            education_info = []
            
            # Look for B.Tech patterns
            btech_patterns = [
                r'B\.?Tech\.?\s*(?:\([^)]*\))?\s*[:\-]?\s*([A-Za-z\s&]+(?:College|University|Institute))',
                r'B\.?Tech\.?\s*(?:\([^)]*\))?\s*([A-Za-z\s&]+(?:College|University|Institute))',
                r'(?:B\.?Tech\.?.*?)([A-Za-z\s&]+(?:Engineering|College|University|Institute))',
            ]
            
            for pattern in btech_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    institution = match.group(1).strip()
                    if len(institution) > 3:  # Filter out too short matches
                        education_info.append(f"B.Tech Institution: {institution}")
            
            # Look for year patterns near B.Tech
            year_patterns = r'B\.?Tech\.?.*?(?:20\d{2}|19\d{2})'
            year_matches = re.finditer(year_patterns, text, re.IGNORECASE)
            for match in year_matches:
                education_info.append(f"B.Tech Year Info: {match.group()}")
            
            # Look for percentage patterns specifically near education keywords (refined)
            percentage_patterns_education = r'(?:B\.?Tech\.?|Diploma|H\.?S\.?C\.?E\.?|S\.?S\.?C\.?E\.?).*?(\d{1,2}(?:.\d{1,2})?)\s*%?' # Look for 1-2 digits, optional decimal, optional %
            percentage_matches_education = re.finditer(percentage_patterns_education, text, re.IGNORECASE)
            for match in percentage_matches_education:
                percentage = match.group(1).strip()
                if percentage:
                    # Basic validation: ensure percentage is reasonable (e.g., <= 100)
                    try:
                        percentage_value = float(percentage)
                        if 0 <= percentage_value <= 100:
                            education_info.append(f"Education Percentage Found (Refined): {percentage}%")
                        else:
                            print(f"Warning: Extracted percentage {percentage}% seems unreasonable.")
                    except ValueError:
                        print(f"Warning: Could not convert extracted percentage {percentage} to float.")
            
            # Broader percentage pattern as a fallback (also refined)
            percentage_patterns_general = r'\b(\d{1,2}(?:.\d{1,2})?)\s*%?' # Look for 1-2 digits, optional decimal, optional %
            percentage_matches_general = re.finditer(percentage_patterns_general, text)
            for match in percentage_matches_general:
                percentage = match.group(1).strip()
                if percentage:
                    # Basic validation: ensure percentage is reasonable (e.g., <= 100)
                    try:
                        percentage_value = float(percentage)
                        if 0 <= percentage_value <= 100 and f"Education Percentage Found (Refined): {percentage}%" not in education_info:
                            education_info.append(f"Percentage Found (General, Refined): {percentage}%")
                    except ValueError:
                        print(f"Warning: Could not convert extracted percentage {percentage} to float (General).")
            
            return '\n'.join(education_info) if education_info else ""
            
        except Exception as e:
            print(f"Error in pattern extraction: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text using pdfplumber for PDF files with table support"""
        try:
            text = ""
            tables_text = ""
            
            with pdfplumber.open(pdf_path) as pdf:
                print(f"Processing {len(pdf.pages)} pages...")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract regular text
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n{page_text}"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        tables_text += f"\n--- Tables from Page {page_num} ---\n"
                        for table_num, table in enumerate(tables, 1):
                            tables_text += f"\nTable {table_num}:\n"
                            # Convert table to readable format
                            if table:
                                df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                                tables_text += df.to_string(index=False) + "\n"
                    
                    print(f"Processed page {page_num}")
            
            # Combine regular text and tables
            combined_text = text
            if tables_text:
                combined_text += "\n\n=== EXTRACTED TABLES ===\n" + tables_text
                
            return combined_text.strip()
        
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None
    
    def extract_text_from_image_standard(self, image_path):
        """Standard image text extraction"""
        try:
            print("Using standard OCR extraction...")
            
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            extracted_text = pytesseract.image_to_string(image, config='--psm 6')
            
            if extracted_text.strip():
                return extracted_text.strip()
            else:
                return "No readable text found in the image."
                
        except Exception as e:
            print(f"Error reading image: {e}")
            return None
    
    def extract_text_from_image(self, image_path):
        """Main image text extraction method"""
        return self.extract_text_with_enhanced_table_support(image_path)
    
    def extract_text_from_docx(self, docx_path):
        """Extract text from DOC/DOCX files"""
        try:
            text = ""
            doc = docx.Document(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables from docx
            tables_text = ""
            for table_num, table in enumerate(doc.tables, 1):
                tables_text += f"\nTable {table_num}:\n"
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data[0] else None)
                    tables_text += df.to_string(index=False) + "\n"
            
            combined_text = text
            if tables_text:
                combined_text += "\n\n=== EXTRACTED TABLES ===\n" + tables_text
                
            return combined_text.strip()
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None
    
    def load_file(self, file_path):
        """Load PDF, DOC/DOCX or image file and prepare for Q&A"""
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found!")
            return False
        
        # Determine file type
       
        file_extension = os.path.splitext(file_path.lower())[1]
        
        print(f"Loading file: {os.path.basename(file_path)}")
        
        if file_extension == '.pdf':
            self.content = self.extract_text_from_pdf(file_path)
            file_type = "PDF"
        elif file_extension in ['.doc', '.docx']:
            self.content = self.extract_text_from_docx(file_path)
            file_type = "DOC/DOCX"
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
            self.content = self.extract_text_from_image(file_path)
            file_type = "Image"
        else:
            print(f"Unsupported file format: {file_extension}")
            print("Supported formats: PDF, DOC, DOCX, JPG, JPEG, PNG, BMP, TIFF, WEBP")
            return False
        
        if not self.content:
            print(f"Failed to extract text from {file_type}!")
            return False
        
        # Show preview of extracted content for debugging
        print("\n=== EXTRACTED CONTENT PREVIEW ===")
        print(self.content[:1000] + "..." if len(self.content) > 1000 else self.content)
        print("=== END PREVIEW ===\n")
        
        # Truncate content if too long
        if len(self.content) > 50000:
            self.content = self.content[:50000] + "\n[Content truncated due to length...]"
            print("Note: Content was truncated due to length limits.")
        
        self.current_file = os.path.basename(file_path)
        
        # Enhanced context prompt for better table understanding
        context_prompt = f"""
        You are an AI assistant that answers questions based on document content, including structured data like tables and forms.
        
        IMPORTANT INSTRUCTIONS:
        - Give direct, accurate answers based on the extracted content
        - Pay special attention to tabular data and structured information
        - When asked about educational qualifications, look for institution names, years, and percentages
        - For B.Tech questions, specifically look for university/college names in the education section
        - Be specific with names, dates, and numbers when available
        - If information is in a table format, extract the exact values
        
        DOCUMENT CONTENT:
        {self.content}
        
        The document has been processed with enhanced table extraction. Answer questions accurately based on this content.
        """
        
        try:
            self.chat = self.model.start_chat()
            response = self.chat.send_message(context_prompt)
            print(f"✅ {file_type} loaded successfully! Enhanced table extraction completed.")
            return True
        except Exception as e:
            print(f"Error initializing chat: {e}")
            return False
    
    def ask_question(self, question):
        """Ask a question with enhanced context understanding"""
        if not self.chat:
            return "Please load a file first using 'load <file_path>' command."
        
        try:
            # Enhanced question processing
            base_question_prompt = f"""
Question: {question}

Please provide a direct, specific answer based on the document content. The document has been analyzed for:
- Education details (degrees, institutions, years, scores)
- Contact information (email, phone, address)
- Skills and qualifications
- Professional experience
- Projects and achievements
- Other relevant sections

Please focus on the specific information requested in the question.
"""

            response = self.chat.send_message(base_question_prompt)
            return response.text

        except Exception as e:
            return f"Error: {e}"
    
    def show_content_preview(self):
        """Show a preview of the extracted content"""
        if not self.content:
            print("No content loaded.")
            return
        
        preview = self.content[:1000] + "..." if len(self.content) > 1000 else self.content
        print(f"\n--- Content Preview from {self.current_file} ---")
        print(preview)
        print("--- End Preview ---\n")

# Initialize chatbot with API key from environment variable
API_KEY = os.getenv('GOOGLE_API_KEY')
if not API_KEY:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")

chatbot = EnhancedMultiFormatChatbot(API_KEY)

@app.route('/')
# def home():
#     return render_template('index.html')
def index():
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"logs/chat_{now}.txt"
    session['chat_log_path'] = filename

    full_path = os.path.join(os.getcwd(), filename)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    return render_template("index.html")

@app.route('/avatar')
def avatar():
    return render_template("avatar.html")

@app.route('/upload', methods=['POST'])
def upload_files():
    print("[DEBUG] Entered /upload route")
    if 'file' not in request.files:
        print("[ERROR] No file part in request")
        return jsonify({'error': 'No file part'}), 400
    
    files = request.files.getlist('file')
    print(f"[DEBUG] Files received: {[file.filename for file in files]}")
    
    if not files or all(file.filename=='' for file in files):
        return jsonify({'error': 'no files selected'}), 400
    

    # if files.filename == "":
    #     return jsonify({'error': 'No selected file'}), 400
    uploaded_files=[]
    session['uploaded_files']=[]

    combined_text=""
    # success_files=[]
    # failed_files=[]
    
    print("DEBUG allowed_file is:", allowed_file)


    for file in files:
        if allowed_file(file.filename):  
            filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[1]
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            if 'uploaded_files' not in session:
                session['uploaded_files']=[]
    
            session['uploaded_files'].append(filepath)
            uploaded_files.append(filename)

            load_success = chatbot.load_file(filepath)

            combined_text += extract_text_from_pdf(filepath) +"\n\n"

            chatbot.load_combined_text(combined_text)
            session['current_file'] = 'multiple_files_combined'
        else:
            print(f"Skipping unsupported file: {file.filename}")

    # if combined_text:
    #     chatbot.load_combined_text(combined_text)
    #     session['current_file'] = 'multiple_files_combined'

    if uploaded_files:
        chatbot.load_combined_text(combined_text)
        session['current_file'] = 'multiple_files_combined'
        return jsonify({
            'success': True, 
            'message': f'{len(uploaded_files)} files uploaded and processed successfully', 
            'files': uploaded_files,
            'ready':True
            }), 200
    else:
        return jsonify({'error': 'No valid files uploaded'}), 400


@app.route('/ask', methods=['POST'])
def ask_question():
    print("[DEBUG] Reached /ask endpoint")
    data = request.get_json()
    print("[DEBUG] Raw JSON received:", data)

    question = data.get('question')
    mute=data.get('mute',False)

    # print(f"received : {question} ,mute: {mute}")

    if not question:
        return jsonify({'error': 'No question provided'}), 400
    
    upload_files=session.get('uploaded_files',[])

    if not upload_files:
        return jsonify({'error': 'Please upload a file first'}), 400
    
    # --- Start: Restore chatbot state from session if needed ---
    combined_text=""
    
    for filepath in upload_files:
        try:
            combined_text += extract_text_from_pdf(filepath) + "\n\n"
        except Exception as e:
            print(f"[ERROR] Failed to extract text from {filepath}: {e}")
    prompt = f"{combined_text}\n\nQuestion: {question}"

    try:
        response_text = chatbot.ask_question(question)
        print("[DEBUG] Got response:", response_text)

        log_conversation(question, response_text)

        if not mute:
            audio_url= generate_polly_audio(response_text)
        else:
            audio_url=None
        return jsonify({
            'response': response_text, 
            'audio_url':audio_url
            })
    
    except Exception as e:
        print("[ERROR] Exception occurred while generating response:", str(e))
        return jsonify({'error': str(e)}), 500

@app.route('/upload_log', methods=['POST'])
def upload_log():
    uploaded_file=request.files.get('logfile')
    override = request.form.get('override') == 'true'
    if uploaded_file:
        mode = 'w' if override else 'a'
        with open(CHAT_LOG_PATH, mode, encoding='utf-8') as f:
            f.write(uploaded_file.read().decode('utf-8'))
        return redirect('/')
    

@app.route('/test-pdf')
def test_pdf():
    text = extract_text_from_pdf('your/test/path/file.pdf')
    return jsonify({'text': text})

@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory('static', filename)

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == "__main__":
    # Check required dependencies
    try:
        pytesseract.get_tesseract_version()
        print("✅ OCR support available")
    except:
        print("⚠️  OCR not available - install Tesseract OCR for image support")
        print("Download from: https://github.com/tesseract-ocr/tesseract")
    
    try:
        import cv2
        print("✅ Enhanced table detection available")
    except ImportError:
        print("⚠️  OpenCV not installed - enhanced table detection not available")
        print("Install with: pip install opencv-python")
    
    # Run the app
    port = int(os.getenv('PORT', 5002))
    app.run(host='0.0.0.0', port=port, debug=True)